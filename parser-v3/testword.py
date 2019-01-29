from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

document = Document()
# 新增新段落
p = document.add_heading('TMD-14921', 0)
# 設定段落對齊方式
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
table = document.add_table(rows=3, cols=2,style='Light List')
table.font = 20
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'New Keyword'
hdr_cells[1].text = 'Reuse Keyword'


hdr_cells = table.rows[1].cells
hdr_cells[0].text = '1'
hdr_cells[1].text = '21'


hdr_cells = table.rows[2].cells
hdr_cells[0].text = '2'
hdr_cells[1].text = '43'

document.save("./writeResult.docx")