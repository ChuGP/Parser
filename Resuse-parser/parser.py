from bs4 import BeautifulSoup
import re
from docx.shared import Pt
import sys
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def set_graph(document,temp):
  p = document.add_heading(test_case_name, 0)
  # 設定段落對齊方式
  p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  paragraph = document.add_paragraph()
  run = paragraph.add_run( temp )
  # run = paragraph.add_run( "New Keyword        Reuse Keyword" )
  font = run.font
  # 設置字體樣式 
  font.name = ' Calibri ' 
  # 設置字體大小 
  font.size = Pt(30)
  font.italic = True

def set_table(document,reuse_keyword):
  size = len(reuse_keyword)
  table = document.add_table(rows=(size+1), cols=1,style='Normal Table')
  i = 0
  for item in reuse_keyword:
    hdr_cells = table.rows[i].cells
    hdr_cells[0].text = "                                " + str(item)
    i+=1
  document.save("./"+ test_case_name +".docx")

if __name__ == '__main__':
  print("分析中請稍後...")
  start = time.time()
  path = './' + sys.argv[1] + '.xml'
  test_case_name = sys.argv[2] #要分析的TMD-XXX
  htmlfile = open(path, 'r', encoding='utf-8')
  htmlhandle = htmlfile.read()
  original = BeautifulSoup(htmlhandle,'xml')
  #抓TEST CASE NAME
  temp = original.findAll('kw',attrs={'type':False,'library':False}) + original.findAll('kw',attrs={'library':['DCTLibrary','Keywords']})
  keywords_list=[ content['name'] for content in temp]
  # deal_keyword = set(keywords_list)
  reuse_keyword = [content for content in set(keywords_list) if '::' in content ]
  reuse_keyword.sort()
  print (reuse_keyword)   #把keyword加進list 
  document = Document()
  # # 新增新段落
  set_graph(document,"                  Reuse Keyword")
  # # 新增表格
  set_table(document,reuse_keyword)
  end = time.time()
  print ("耗時 : ", end - start, "seconds.")