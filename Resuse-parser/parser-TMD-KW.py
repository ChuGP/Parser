from bs4 import BeautifulSoup
import re
from docx.shared import Pt
import sys
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def addlist(package):
  keywords_list=[ content['name'] for content in package]
  return set(keywords_list)   #把keyword加進list 

def namedict(contents):
  name_dict={}
  for content in contents:
     if "TMD" in content['name']:name_dict[content['name'].split(' ', 1 )[0]] = content['name']
  return name_dict

def keyword_dict(name_dict,original): 
  Keyword_save = {}
  for item in name_dict:
    temp = original.find('suite',{'name':name_dict[item]}) 
    package = temp.findAll('kw',attrs={'library':['DCTLibrary','Keywords']}) + temp.findAll('kw',attrs={'type':False,'library':False})
    Keyword_save[item] = addlist(package)
  return Keyword_save

def get_master_keyword(Keyword_save,test_case_name):
  default=set()
  for item in Keyword_save: 
    if ( item != test_case_name ): default = ( default | Keyword_save[item] )
  return default

def set_graph(document,temp):
  p = document.add_heading(test_case_name, 0)
  # 設定段落對齊方式
  p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  paragraph = document.add_paragraph()
  run = paragraph.add_run( temp )
  # run = paragraph.add_run( "New Keyword        Reuse Keyword" )
  font = run.font
  # 設置字體樣式 
  font.name = ' Calibri ' 
  # 設置字體大小 
  font.size = Pt(30)
  font.italic = True

def set_table(document,reuse_keywords):
  size = len(reuse_keywords)
  table = document.add_table(rows=(size+1), cols=1,style='Normal Table')
  hdr_cells = table.rows[0].cells 
  i = 0
  for item in reuse_keywords:
    hdr_cells = table.rows[i].cells
    hdr_cells[0].text = str(item)
    i+=1
  document.save("./"+ test_case_name +".docx")

if __name__ == '__main__':
  print("分析中請稍後...")
  start = time.time()
  path = './' + sys.argv[1] + '.xml'
  test_case_name = sys.argv[2] #要分析的TMD-XXX
  htmlfile = open(path, 'r', encoding='utf-8')
  htmlhandle = htmlfile.read()
  original = BeautifulSoup(htmlhandle,'xml')
  #抓TEST CASE NAME
  contents = original.findAll('suite',attrs={'name':True})
  name_dict = namedict(contents) #做名字的dict key為縮寫 value為全名
  Keyword_save = keyword_dict(name_dict,original) #做keyword的dict key為TMD-XXX value用到的keyword list

  Keyword_expand_list =[]
  for item in name_dict:
    Keyword_expand_list += Keyword_save[item]
  
  # print(Keyword_expand_list)
  Keyword_expand = {}
  ans = []
  for item in Keyword_expand_list:
    Tmd_name = []
    for name in name_dict:
      if item in Keyword_save[name]:  
        Tmd_name.append(name)

    Keyword_expand[item] = Tmd_name
    token = item + ":\n" + ",  ".join(Tmd_name) + "\n"
    ans.append(token)

  document = Document()
  # 新增新段落
  set_graph(document,"Reuse Keywords And Distribute")
  # 新增表格
  set_table(document,ans)
  end = time.time()
  print ("耗時 : ", end - start, "seconds.")